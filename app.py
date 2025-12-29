import os
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from PIL import Image as PILImage
from agno.agent import Agent
from agno.models.google import Gemini
from agno.media import Image as AgnoImage
from docx import Document
from docx.shared import Inches
import io
import re
from dotenv import load_dotenv
import pydicom
import numpy as np
import tempfile
import base64
import time
import gc

# Load environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

app = Flask(__name__)
CORS(app)

# Medical Analysis Query
MEDICAL_QUERY = """
You are a highly skilled medical expert specializing in diagnostic imaging and analytical interpretation of medical graphs. Analyze the provided medical image or graph comprehensively, solely based on the visual data available, and structure your analysis clearly under the following headings:

### 1. Image/Graph Type & Anatomical Region
- Specify the imaging or graph modality clearly (X-ray, MRI, CT, Ultrasound, ECG, EEG, etc.)
- Identify the anatomical region or physiological system depicted
- Comment on the technical quality, clarity, and adequacy of the data provided


### 2. Detailed Observations & Key Findings
- Systematically list primary visual or numerical observations
- Clearly describe any abnormalities, anomalies, or deviations from normal reference ranges
- Include precise measurements, numerical values, or densities as applicable
- Clearly describe the location, size, shape, characteristics, or pattern of abnormalities
- Rate severity clearly as Normal, Mild, Moderate, or Severe

### 3. Diagnostic Assessment
- State your primary diagnosis clearly and confidently based solely on visual or graphical evidence
- List possible differential diagnoses ranked by likelihood, supported by the provided data
- Highlight specific visual or numerical evidence underpinning each diagnosis
- Clearly flag any critical, urgent, or emergent findings requiring immediate medical attention

### 4. Patient-Friendly Explanation
- Clearly and simply explain your findings without medical jargon
- Provide definitions or analogies to help the patient understand the significance
- Address potential patient concerns about severity, prognosis, or immediate next steps
- Suggest general recommendations or follow-up steps clearly understandable by a non-medical individual

Ensure your analysis remains precise, thorough, and clear, making it fully applicable to multimodal medical data including both images and graphical presentations. Consider that your interpretation may need to stand independently without additional patient history or symptomatic context.
"""

# Initialize medical agent
medical_agent = None
if GOOGLE_API_KEY:
    medical_agent = Agent(
        model=Gemini(
            id="gemini-2.0-flash",
            api_key=GOOGLE_API_KEY
        ),
        tools=[],
        markdown=True
    )


def process_image(file):
    """Process uploaded image (standard or DICOM) and return PIL Image"""
    file_name = file.filename.lower()
    
    if file_name.endswith(".dcm") or file_name.endswith(".dicom"):
        try:
            # Reset file pointer to beginning
            file.seek(0)
            ds = pydicom.dcmread(file)
            pixel_array = ds.pixel_array.astype(float)
            # Normalize to 0-255 for display
            pixel_array -= pixel_array.min()
            if pixel_array.max() > 0:
                pixel_array /= pixel_array.max()
            pixel_array = (pixel_array * 255).astype("uint8")
            image = PILImage.fromarray(pixel_array)
        except Exception as e:
            raise ValueError(f"Failed to read DICOM file: {e}")
    else:
        # Reset file pointer to beginning
        file.seek(0)
        image = PILImage.open(file)
    
    return image


def resize_image(image, max_width=500):
    """Resize image maintaining aspect ratio"""
    width, height = image.size
    aspect_ratio = width / height
    new_width = max_width
    new_height = int(new_width / aspect_ratio)
    return image.resize((new_width, new_height))


def create_docx(analysis_markdown, image_pil):
    """Create DOCX document from analysis and image"""
    doc = Document()
    doc.add_heading('Medical Imaging Analysis', 0)
    # Add image at the top
    doc.add_heading('Uploaded Image', level=1)
    img_stream = io.BytesIO()
    image_pil.save(img_stream, format='PNG')
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(5))
    # Add analysis after image
    bold_pattern = re.compile(r'(\*\*|__)(.+?)\1')
    for line in analysis_markdown.split('\n'):
        if line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        elif line.strip() == '---':
            doc.add_page_break()
        elif line.strip():
            # Handle bold markdown
            p = doc.add_paragraph()
            last_idx = 0
            for m in bold_pattern.finditer(line):
                # Add text before bold
                if m.start() > last_idx:
                    p.add_run(line[last_idx:m.start()])
                # Add bold text
                p.add_run(m.group(2)).bold = True
                last_idx = m.end()
            # Add any remaining text
            if last_idx < len(line):
                p.add_run(line[last_idx:])
    return doc


@app.route('/')
def index():
    """Serve the main page"""
    return render_template('index.html', api_configured=bool(GOOGLE_API_KEY))


@app.route('/api/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({
        'status': 'ok',
        'api_configured': bool(GOOGLE_API_KEY)
    })


@app.route('/api/upload', methods=['POST'])
def upload_image():
    """Handle image upload and return preview"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Process image
        image = process_image(file)
        resized_image = resize_image(image)
        
        # Convert to base64 for preview
        img_buffer = io.BytesIO()
        resized_image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
        
        return jsonify({
            'success': True,
            'preview': f'data:image/png;base64,{img_base64}'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/analyze', methods=['POST'])
def analyze_image():
    """Analyze uploaded image"""
    try:
        if not medical_agent:
            return jsonify({'error': 'Medical agent not configured. Please set GOOGLE_API_KEY.'}), 500
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Process image
        image = process_image(file)
        resized_image = resize_image(image)
        
        # Save to temp file for analysis
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_path = temp_file.name
        temp_file.close()  # Close the file handle immediately
        
        agno_image = None
        try:
            resized_image.save(temp_path)
            
            # Create AgnoImage object
            agno_image = AgnoImage(filepath=temp_path)
            
            # Run analysis
            response = medical_agent.run(MEDICAL_QUERY, images=[agno_image])
            
            # Convert image to base64 for storage
            img_buffer = io.BytesIO()
            resized_image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
            
            return jsonify({
                'success': True,
                'analysis': response.content,
                'image_base64': f'data:image/png;base64,{img_base64}'
            })
        finally:
            # Clean up temp file - ensure it's deleted even if there's an error
            # Release any references first
            if agno_image is not None:
                try:
                    del agno_image
                except:
                    pass
            gc.collect()  # Force garbage collection
            
            # Retry mechanism for Windows file locking
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)
                    break  # Success, exit retry loop
                except (OSError, PermissionError) as e:
                    if attempt < max_retries - 1:
                        time.sleep(0.1)  # Wait 100ms before retry
                    else:
                        # On final attempt, schedule deletion on exit if still locked
                        try:
                            import atexit
                            def cleanup_on_exit(path=temp_path):
                                try:
                                    if os.path.exists(path):
                                        os.unlink(path)
                                except:
                                    pass
                            atexit.register(cleanup_on_exit)
                        except:
                            pass
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    """Generate and download DOCX file"""
    try:
        data = request.json
        if not data or 'analysis' not in data or 'image_base64' not in data:
            return jsonify({'error': 'Missing analysis or image data'}), 400
        
        # Decode image from base64
        image_data = data['image_base64'].split(',')[1]
        image_bytes = base64.b64decode(image_data)
        image_pil = PILImage.open(io.BytesIO(image_bytes))
        
        # Create DOCX
        doc = create_docx(data['analysis'], image_pil)
        
        # Save to BytesIO
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        return send_file(
            docx_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='medical_image_analysis.docx'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)

